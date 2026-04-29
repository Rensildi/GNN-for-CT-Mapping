"""Convert the markdown report into a .docx via python-docx.

Pandoc is the right tool for this kind of conversion, but it requires
sudo to install on this machine. python-docx is already in the venv,
so we hand-roll a small converter that handles the markdown features
this report actually uses:

  - ATX headings (#, ##, ###, ####)
  - Paragraphs with inline bold (**text**) and italic (*text*)
  - Bullet lists (-, *)
  - Numbered lists (1. 2. 3.)
  - GitHub-style tables (| col | col |)
  - Horizontal rules (---)
  - Block-level code fences (```...```) — rendered as monospace paragraphs
  - Italic-line emphasis (*single-line italic blocks*)

This is intentionally not a full markdown parser. It's tuned to the
specific shape of `gnn_lung_nodule_classification.md` and produces a
clean, readable Word document for delivery.

Run:
    source AI/bin/activate
    python report/build_report_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent
MD_PATH = REPO_ROOT / "report" / "gnn_lung_nodule_classification.md"
DOCX_PATH = REPO_ROOT / "report" / "gnn_lung_nodule_classification.docx"


# Inline bold/italic regexes. Bold is matched first so that
# `**word**` doesn't also trip the italic regex.
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")

# Image syntax: ![alt text](path). For the docx output we don't embed the
# image; we render a centered, italic placeholder paragraph carrying the
# alt text so the caller can drop the actual figure in by hand.
IMAGE_RE = re.compile(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$")

# Document body font. Times New Roman is the requested research-paper face.
BODY_FONT = "Times New Roman"
MONO_FONT = "Courier New"


def _add_inline_runs(paragraph, text: str, base_bold: bool = False, base_italic: bool = False) -> None:
    """Tokenize text into inline-formatting runs and append to a paragraph.

    Bold > italic > code > plain. We tokenize by repeatedly finding the
    earliest matching span and emitting a plain run before it + a styled
    run for it.
    """
    cursor = 0
    while cursor < len(text):
        # Find the earliest of bold / italic / inline-code matches starting
        # at or after `cursor`.
        candidates = []
        for kind, regex in [("bold", BOLD_RE), ("italic", ITALIC_RE), ("code", INLINE_CODE_RE)]:
            m = regex.search(text, cursor)
            if m:
                candidates.append((m.start(), kind, m))
        if not candidates:
            # No more inline formatting; emit the rest as plain text.
            run = paragraph.add_run(text[cursor:])
            run.bold = base_bold
            run.italic = base_italic
            break
        candidates.sort(key=lambda t: t[0])
        start, kind, m = candidates[0]
        if start > cursor:
            run = paragraph.add_run(text[cursor:start])
            run.bold = base_bold
            run.italic = base_italic
        styled = paragraph.add_run(m.group(1))
        if kind == "bold":
            styled.bold = True
            styled.italic = base_italic
        elif kind == "italic":
            styled.italic = True
            styled.bold = base_bold
        elif kind == "code":
            styled.font.name = MONO_FONT
            styled.bold = base_bold
            styled.italic = base_italic
        cursor = m.end()


def _is_table_separator(line: str) -> bool:
    """Markdown table separator row, e.g. '|:---|:---:|---:|'.

    Strict CommonMark requires three or more dashes per cell, but real-world
    markdown tables sometimes use shorter forms like ':--:' or ':-:' for
    narrow columns. We accept any positive run of dashes with optional
    leading/trailing alignment colons.
    """
    s = line.strip()
    if not (s.startswith("|") and s.endswith("|")):
        return False
    inner = s[1:-1]
    cells = [c.strip() for c in inner.split("|")]
    if not cells:
        return False
    return all(re.fullmatch(r":?-+:?", c) for c in cells)


def _split_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _emit_paragraph(doc, text: str, *, italic: bool = False, style: str = "Normal") -> None:
    """Add a regular paragraph with inline-formatted runs."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    _add_inline_runs(p, text, base_italic=italic)


def _emit_table(doc, rows: list[list[str]]) -> None:
    """Render a markdown table as a python-docx Table."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    # Pad short rows so column count is uniform.
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            # Replace any existing default paragraph with our own.
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            _add_inline_runs(p, cell_text, base_bold=(i == 0))
    # Compact spacing
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()  # spacer after the table


def _emit_code_block(doc, code: str) -> None:
    """Monospace block. Multiple-line code is concatenated with line breaks."""
    p = doc.add_paragraph()
    run = p.add_run(code.rstrip())
    run.font.name = MONO_FONT
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def _emit_figure_placeholder(doc, alt: str, path: str) -> None:
    """Render an image reference as a centered placeholder paragraph.

    The user wants the .md file to embed real images but the .docx file
    to leave placeholders so that figures can be dropped in by hand at
    layout time.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    label = alt.strip() if alt.strip() else Path(path).name
    run = p.add_run(f"[Figure placeholder: {label}]")
    run.italic = True
    run.font.name = BODY_FONT
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _setup_styles(doc) -> None:
    """Tune base paragraph + heading styles for a research-paper look in
    Times New Roman.
    """
    styles = doc.styles
    # Base body style.
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    # Heading sizes.
    for name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 12), ("Heading 4", 11)]:
        try:
            s = styles[name]
            s.font.name = BODY_FONT
            s.font.size = Pt(size)
            s.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            s.font.bold = True
        except KeyError:
            pass
    # List paragraph styles also need the body font set explicitly so they
    # don't fall back to the docx default.
    for name in ("List Bullet", "List Number"):
        try:
            styles[name].font.name = BODY_FONT
        except KeyError:
            pass


def convert(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    _setup_styles(doc)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Empty line — skip; paragraph spacing handled in style.
        if not stripped:
            i += 1
            continue

        # Horizontal rule: skip silently. The user does not want line break
        # separators in the docx; section structure is carried by headings.
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # Headings
        h_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if h_match:
            level = len(h_match.group(1))
            heading_text = h_match.group(2).strip()
            # H1 in markdown -> document title (large, centered).
            if level == 1:
                p = doc.add_heading("", level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(heading_text)
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            else:
                doc.add_heading(heading_text, level=level - 1)
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            _emit_code_block(doc, "\n".join(buf))
            continue

        # Table — header row + separator + body rows
        if stripped.startswith("|") and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            header = _split_table_row(lines[i])
            i += 2
            body_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                body_rows.append(_split_table_row(lines[i]))
                i += 1
            _emit_table(doc, [header] + body_rows)
            continue

        # Bulleted list
        if re.match(r"^[-*]\s+", stripped):
            while i < len(lines):
                m = re.match(r"^([-*])\s+(.*)$", lines[i].strip())
                if not m:
                    break
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                _add_inline_runs(p, m.group(2))
                i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines):
                m = re.match(r"^\d+\.\s+(.*)$", lines[i].strip())
                if not m:
                    break
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(2)
                _add_inline_runs(p, m.group(1))
                i += 1
            continue

        # Image reference. The .md embeds real images; the .docx renders
        # only a placeholder so figures can be inserted at layout time.
        img_match = IMAGE_RE.match(line)
        if img_match:
            _emit_figure_placeholder(doc, img_match.group(1), img_match.group(2))
            i += 1
            continue

        # Italic-only paragraph (closing footer line).
        if stripped.startswith("*") and stripped.endswith("*") and stripped.count("*") == 2:
            _emit_paragraph(doc, stripped[1:-1], italic=True)
            i += 1
            continue

        # Regular paragraph — accumulate consecutive non-empty lines.
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            # Stop at any block-level start.
            if (re.match(r"^#{1,6}\s+", nxt)
                    or nxt in ("---", "***", "___")
                    or nxt.startswith("```")
                    or re.match(r"^[-*]\s+", nxt)
                    or re.match(r"^\d+\.\s+", nxt)
                    or (nxt.startswith("|") and i + 1 < len(lines)
                        and _is_table_separator(lines[i + 1]))):
                break
            para_lines.append(lines[i])
            i += 1
        merged = " ".join(l.strip() for l in para_lines)
        _emit_paragraph(doc, merged)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    convert(MD_PATH, DOCX_PATH)
